import abc
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from songbook2docx.utils.style_manager import StyleManager


class StyledCell(metaclass=abc.ABCMeta):
    @classmethod
    def __subclasshook__(cls, subclass):
        return (hasattr(subclass, 'add_runs_to_paragraph') and
                callable(subclass.add_runs_to_paragraph) or
                NotImplemented)

    @abc.abstractmethod
    def add_runs_to_paragraph(self, par: Paragraph, style_manager: StyleManager) -> list[Run]:
        raise NotImplementedError
