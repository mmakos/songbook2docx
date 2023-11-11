import re

from docx.styles.style import BaseStyle
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from songbook2docx.styled.chord import Chord
from songbook2docx.styled.chord import HIDE_KEY_MARK_FLAG
from songbook2docx.styled.styled_cell import StyledCell
from songbook2docx.utils.style_manager import StyleManager, KEY, CHORDS
from songbook2docx.utils.styler import Styler

chord_delimiters = (
    r"(?<!<)/",  # Oddziela slash, pod warunkiem, że nie jest to tag html </
    r" ?\(\^",  # Oddziela podwyższenie tonacji z ewentualną spacją (^
    r" \(|\) ",  # Oddziela nawias początkowy lub końcowy ze spacją
    r"(?<!>)\(",  # Oddziela nawias początkowy, pod warunkiem, że nie ma przed nim końca tagu (wtedy jest to opcjonalny bas/interwał)
    r"\)(?!<)",  # Oddziela nawias końcowy, pod warunkiem, że nie ma za nim tagu html (wtedy jest to koniec opcjonalnego basu/interwału)
    r"\|x[0-9]+",  # Oddziela repetycję akordów
    r"[ |]+(?![0-9])",  # Oddziela spacje i pionową kreskę, pod warunkiem, że za nią nie ma cyfr (wtedy jest to kolejny interwał, np. A2-1 4-3)
    r"… ?"  # Oddziela wielokropek (zapętlone akordy)
)


class Chords(StyledCell):
    def __init__(self, chords: list[Chord], key: str | None, styles: int):
        self.chords = chords
        self.key = key
        self.styles = styles

    @staticmethod
    def chords_from_html(html: str):
        chords, key, styles = Chords.__parse_html(html)
        return Chords(chords, key, styles)

    @staticmethod
    def __parse_html(html: str) -> tuple[list[Chord], str | None, int]:
        key = None
        chords_start = max(html.find("""<b class="chord">"""), 0)

        if chords_start > 0:
            key = html[:chords_start].strip()
            if not key.strip().startswith("<i><b>"):
                key = None
            else:
                key = key.replace("<i><b>", "")
                key = key.replace("</b></i>", "")
        html = html[chords_start:]

        html = html.replace("""<b class="chord">""", "").replace("</b>", "")
        html, styles = Chords.__parse_styles(html)
        chords_str: list[tuple[str, str]] = split_chords(f"({'|'.join(chord_delimiters)})", html)
        chords = list()
        for chord_str in chords_str:
            chords.append(Chord.chord_from_text(chord_str[0], chord_str[1]))
        return chords, key, styles

    @staticmethod
    def __parse_styles(html: str) -> tuple[str, int]:
        styles = 0
        if html.startswith("<i>") and html.endswith("</i>"):
            styles |= Styler.ITALIC
            html = html[len("<i>"): -len("</i>")]
        return html, styles

    def transpose(self, interval: int):
        for chord in self.chords:
            chord.transpose(interval)

    def apply_flags(self, flags: int):
        if flags & HIDE_KEY_MARK_FLAG > 0:
            self.key = None
        proceeded_chords: list[Chord] = list()
        for chord in self.chords:
            proceeded_chords.extend(chord.apply_flags(flags))
        distinct_chords: list[Chord] = list()
        for i in range(len(proceeded_chords) - 1):
            if not proceeded_chords[i].is_same_chord(proceeded_chords[i + 1]):
                distinct_chords.append(proceeded_chords[i])
        if len(proceeded_chords) > 0:
            distinct_chords.append(proceeded_chords[-1])
        self.chords = distinct_chords

    def add_runs_to_paragraph(self, par: Paragraph, style_manager: StyleManager) -> list[Run]:
        runs = list()
        if self.key is not None:
            styled_keys = Styler.style_text(self.key)
            for key in styled_keys:
                run = par.add_run(key[0], style_manager.get_style(KEY))
                Styler.style_run(run, key[1])
                runs.append(run)
        if len(self.chords) > 0:
            chord_style = style_manager.get_style(CHORDS)
            if self.key is not None:
                runs.append(par.add_run(" ", chord_style))
            for chord in self.chords:
                runs.extend(self.add_chord_runs(par, chord, chord_style))
                runs.append(par.add_run(chord.delimiter, chord_style))

        return runs

    @staticmethod
    def add_chord_runs(par: Paragraph, chord: Chord, style: BaseStyle) -> list[Run]:
        runs = list()
        if not chord.chord:
            return list()
        runs.append(par.add_run(chord.chord + chord.aug, style))
        if chord.base:
            base = Styler.style_text(chord.base)  # Bo może byc przekreślone <s>1</s>
            for b in base:
                run = par.add_run(b[0], style)
                Styler.style_run(run, b[1] | Styler.SUB)
                runs.append(run)
        if chord.add:
            run = par.add_run(chord.add, style)
            Styler.style_run(run, Styler.SUPER)
            runs.append(run)

        return runs


def split_chords(regex: str, string: str) -> list[tuple[str, str]]:
    """Splits string with regex, removes empty delimeters and aligns to pairs (chord -> delimeter)"""
    if not string:
        return list()
    split = [s for s in re.split(regex, string) if len(s) > 0]
    i = 0
    while i < len(split) - 1:
        if not re.match("^[A-Za-z]", split[i]) and not re.match("^[A-Za-z]", split[i + 1]):
            split[i] = split[i] + split.pop(i + 1)
        else:
            i += 1

    if not re.match("^[A-Za-z]", split[0]):
        split.insert(0, "")
    if len(split) % 2 > 0:
        split.append("")
    return [(split[i], split[i + 1]) for i in range(0, len(split) - 1, 2)]
