import xml.etree.ElementTree as et

from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from songbook2docx.styled.chord import HIDE_ALTERNATIVE_KEY_FLAG
from songbook2docx.styled.chords import Chords
from songbook2docx.styled.repetition import Repetition
from songbook2docx.styled.song_text import SongText
from songbook2docx.styled.styled_cell import StyledCell
from songbook2docx.utils.style_manager import StyleManager, TITLE, AUTHOR, SONG
from songbook2docx.utils.text_width_provider import TextWidthProvider


class Column:
    def __init__(self, strings: list[str], cell_type: int):
        self.strings: list[str] = strings
        self.cell_type: int = cell_type


class Row:
    def __init__(self, cells: list[StyledCell]):
        self.cells: list[StyledCell] = cells


class Song:
    def __init__(self, html: et.Element = None):
        self.title: str = ""
        self.authors: str = ""
        self.flags: int = 0
        self.transposition: int = 0
        self.rows: list[Row]
        if html is not None:
            self.rows: list[Row] = self.__parse_html(html)

    def parse_options(self, html: et.Element):
        self.title = Song.__title_from_html(html)
        self.authors = Song.__authors_from_html(html)
        self.flags = Song.__flags_from_html(html)
        self.transposition = Song.__transposition_from_html(html)

    def __parse_html(self, html: et.Element) -> list[Row]:
        self.parse_options(html)

        table_columns = Song.__columns_from_table([tr for tr in html.iter("tr")][0])
        columns: list[Column] = list()
        for i, col in enumerate(table_columns):
            rows = Song.__rows_from_column(col)
            columns.append(Column(rows, Song.__get_column_type(rows, i)))
        return Song.__transform_columns_to_rows(columns)

    @staticmethod
    def __rows_from_column(column: et.Element) -> list[str]:
        return [et.tostring(m, encoding="unicode").replace("<span>", "")
                .replace("</span>", "")
                .replace("<br />", "")
                .replace("<br>", "")
                .replace("<emsp />", "\t")
                .replace("\t\t", "\t") for m in get_from_tags(column, "span")]

    @staticmethod
    def __columns_from_table(html: et.Element) -> list[et.Element]:
        return get_from_tags(html, "td")

    @staticmethod
    def __title_from_html(html: et.Element) -> str:
        title_strings = get_from_tags(html, "h2")
        if len(title_strings) > 0:
            return title_strings[0].text
        return ""

    @staticmethod
    def __authors_from_html(html: et.Element) -> str:
        divs = html.iter("div")
        for div in divs:
            if div.attrib.get("class") == "author":
                return ", ".join(div.itertext())
        return ""

    @staticmethod
    def __flags_from_html(html: et.Element) -> int:
        flags_strings = get_from_tags(html, "flags")
        if len(flags_strings) > 0:
            try:
                return int(flags_strings[0].text)
            except:
                pass
        return 0

    @staticmethod
    def __transposition_from_html(html: et.Element) -> int:
        transposition_strings = get_from_tags(html, "transposition")
        if len(transposition_strings) > 0:
            try:
                return int(transposition_strings[0].text)
            except:
                pass
        return 0

    @staticmethod
    def __transform_columns_to_rows(columns: list[Column]) -> list[Row]:
        if not columns:
            return list()
        rows = list()
        for i in range(len(columns[0].strings)):
            cells: list[StyledCell] = list()
            for col in columns:
                col_strings = col.strings[i] if i < len(col.strings) else ""
                if col.cell_type == CellType.TEXT or col.cell_type == CellType.OTHER:
                    cells.append(SongText(col_strings))
                elif col.cell_type == CellType.REPETITION:
                    cells.append(Repetition(col_strings))
                elif col.cell_type == CellType.CHORD:
                    cells.append(Chords.chords_from_html(col_strings))
            rows.append(Row(cells))
        return rows

    @staticmethod
    def __get_column_type(rows: list[str], column_idx: int) -> int:
        if column_idx == 0:
            return CellType.TEXT
        for row in rows:
            if 'class="chord"' in row:
                return CellType.CHORD
        for row in rows:
            if '|x' in row:
                return CellType.REPETITION
        return CellType.OTHER

    def add_paragraphs_to_doc(self, content_par: Paragraph, tab_stops_offset: float, show_authors: bool, style_manager: StyleManager,
                              text_width_provider: TextWidthProvider):
        content_par.insert_paragraph_before(self.title, style=style_manager.get_style(TITLE))
        if show_authors and len(self.authors) > 0:
            content_par.insert_paragraph_before(self.authors, style=style_manager.get_style(AUTHOR))
        pars = list()
        # dla każdej kolumny lista szerokości komórek wierszy
        cell_lengths: list[list[float]] = [[0.0 for _ in range(len(self.rows))] for _ in range(len(self.rows[0].cells))]
        for row_index, row in enumerate(self.rows):
            par: Paragraph = content_par.insert_paragraph_before(style=style_manager.get_style(SONG))
            pars.append(par)

            for i, cell in enumerate(row.cells):
                added_runs = cell.add_runs_to_paragraph(par, style_manager)
                width = 0
                if any(run.text for run in added_runs):
                    width = text_width_provider.get_text_width(added_runs) + Cm(tab_stops_offset).pt
                    if i == 0 and par.paragraph_format.left_indent:
                        width += par.paragraph_format.left_indent.pt
                cell_lengths[i][row_index] = width
                if (width > 0 or i == 0) and i < len(row.cells) - 1:
                    tab_run = par.add_run()
                    tab_run.add_tab()

        for i in range(len(cell_lengths) - 1):
            widths = [w for w, z in zip(cell_lengths[i], cell_lengths[i + 1]) if z > 0]
            max_width = 0.0001
            if len(widths) > 0:
                max_width = max(widths)

            tab_stop_added = False
            for j, par in enumerate(pars):
                if cell_lengths[i + 1][j] > 0:
                    tab_stop_added = True
                    par.paragraph_format.tab_stops.add_tab_stop(Pt(max_width))
            if tab_stop_added:
                for j, _ in enumerate(cell_lengths[i + 1]):
                    cell_lengths[i + 1][j] += max(cell_lengths[i][j], max_width)

    def transpose(self):
        for row in self.rows:
            for cell in row.cells:
                if type(cell) == Chords:
                    cell.transpose(self.transposition)

    def apply_flags(self, global_flags: int):
        flags = self.flags | global_flags
        if len(self.rows) <= 0:
            return
        if flags & HIDE_ALTERNATIVE_KEY_FLAG > 0:
            self.__remove_alternative_chords()

        for row in self.rows:
            for cell in row.cells:
                if type(cell) == Chords:
                    cell.apply_flags(flags)

    def __remove_alternative_chords(self):
        chord_lines = [0 if type(cell) == Chords else -1 for cell in self.rows[0].cells]  # lista przechowuje ilość faktycznych linii akordów w każdej kolumnie
        first_chord_column_idx = chord_lines.index(0)
        if len([c for c in chord_lines if c == 0]) > 1:
            for row in self.rows:
                for i, cell in enumerate(row.cells):
                    if type(cell) == Chords:
                        if len(cell.chords) > 0:
                            chord_lines[i] += 1
            for i in range(first_chord_column_idx + 1, len(chord_lines)):
                if chord_lines[i] > 0.9 * chord_lines[first_chord_column_idx]:
                    for row in self.rows:
                        row.cells.pop(i)


def get_from_tags(html: et.Element, tag: str) -> list[et.Element]:
    return html.findall(tag)


class CellType:
    TEXT = 0
    CHORD = 1
    REPETITION = 2
    OTHER = 3
