import copy
import io
import os
import sys
import traceback
import urllib.request
from http.cookies import SimpleCookie
from urllib.parse import parse_qs

from docx import Document

from config import *
from songbook2docx.parser import parse_html
from songbook2docx.song import Song
from songbook2docx.utils.style_manager import StyleManager
from songbook2docx.utils.text_width_provider import TextWidthProvider

sys.path.insert(0, os.path.dirname(__file__))
BATCH_SIZE = 2048
PHP_SCRIPT_URL = "https://spiewnik.mmakos.pl/export_meeting2.php"
FILENAME_PREFIX = "spiewnik"
GLOBAL_CONFIG = parse_config("conf.ini")


def application(environ, start_response):
    try:
        cookie = SimpleCookie()
        cookie.load(environ.get("HTTP_COOKIE", ""))

        meetingId = cookie.get("meeting")
        if meetingId is not None:
            meetingId = meetingId.value
            content = urllib.request.urlopen(f"{PHP_SCRIPT_URL}?meetingId={meetingId}").read()

            query_string: str = environ['QUERY_STRING']
            if query_string.startswith("?"):
                query_string = query_string[1:]
            params = parse_qs(query_string)

            docx: Document = convertToDocx(content, params)

            docxBytesIO = io.BytesIO()
            docx.save(docxBytesIO)

            mv = docxBytesIO.getbuffer()
            length = len(mv)
            start_response('200 OK',
                           [('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                            ('Content-Length', str(length)),
                            ('Content-Disposition', f"attachment; filename={FILENAME_PREFIX}_{meetingId}.docx")])

            return [bytes(mv[i: i + BATCH_SIZE]) for i in range(0, length, BATCH_SIZE)]
        else:
            start_response('200 OK', [('Content-Type', 'text/plain')])
            return [bytes(f'not-in-meeting', "utf-8")]
    except Exception:
        error = traceback.format_exc()
        start_response('200 OK', [('Content-Type', 'text/plain')])
        return [bytes(error, "utf-8")]


def convertToDocx(plain: bytes, params: dict) -> Document:
    doc: Document = Document(TEMPLATE_PATH)
    config = Config(params_to_config(params))

    songs: list[Song] = parse_html(plain.decode("UTF-8"))

    style_manager = StyleManager(doc, config.config)
    text_width_provider = TextWidthProvider("fonts", config.wanted_font, DEFAULT_FONT, config.font_size)

    content_par = [par for par in doc.paragraphs if "SongbookContent" in par.text][0]

    for song in songs:
        song.apply_flags(config.get_chord_flags())
        song.transpose()
        song.add_paragraphs_to_doc(content_par, config.tab_stops_offset, config.show_author, style_manager, text_width_provider)

    content_el = content_par._element
    content_el.getparent().remove(content_el)
    content_el._p = content_el._element = None

    return doc


def params_to_config(params: dict) -> dict:
    config = copy.deepcopy(GLOBAL_CONFIG)
    init_dict_if_absent(config, "general")
    init_dict_if_absent(config, "chord-flags")
    init_dict_if_absent(config, "style-song-content")
    init_dict_if_absent(config, "style-chords")
    init_dict_if_absent(config, "style-repetition")
    init_dict_if_absent(config, "style-key")

    general = config["general"]
    general["show-author"] = convert_boolean(params, "show-author")
    set_if_present(general, params, "tab-stops-offset")

    chord_flags = config["chord-flags"]
    chord_flags["hide-uncommon-added-interval"] = convert_boolean(params, "hide-uncommon-added-interval")
    chord_flags["aug-and-dim-guitar-mode"] = convert_boolean(params, "aug-and-dim-guitar-mode")
    chord_flags["divide-delays"] = convert_boolean(params, "divide-delays")
    chord_flags["hide-incomplete-chords"] = convert_boolean(params, "hide-incomplete-chords")
    chord_flags["simplify-multiply"] = convert_boolean(params, "simplify-multiply")
    chord_flags["simplify-aug-to-guitar"] = convert_boolean(params, "simplify-aug-to-guitar")
    chord_flags["hide-base"] = convert_boolean(params, "hide-base")
    chord_flags["hide-alternative-key-flag"] = convert_boolean(params, "hide-alternative-key-flag")
    chord_flags["hide-key-mark-flag"] = convert_boolean(params, "hide-key-mark-flag")

    style_song = config["style-song-content"]
    set_if_present(style_song, params, "font-size")
    set_if_present(style_song, params, "font-family")
    style_song["font-bold"] = convert_boolean(params, "text-style-bold")
    style_song["font-italic"] = convert_boolean(params, "text-style-italic")

    style_chords = config["style-chords"]
    style_chords["font-bold"] = convert_boolean(params, "chord-style-bold")
    style_chords["font-italic"] = convert_boolean(params, "chord-style-italic")

    style_repetition = config["style-repetition"]
    style_repetition["font-bold"] = convert_boolean(params, "repetition-style-bold")
    style_repetition["font-italic"] = convert_boolean(params, "repetition-style-italic")

    style_key = config["style-key"]
    style_key["font-bold"] = convert_boolean(params, "tonation-style-bold")
    style_key["font-italic"] = convert_boolean(params, "tonation-style-italic")

    return config


def convert_boolean(params: dict, key: str) -> str:
    if key not in params:
        return 'no'
    value = params[key]
    return 'yes' if len(value) > 0 and value[0] == 'on' else 'no'


def set_if_present(config: dict, params: dict, key: str):
    if key in params:
        value = params[key]
        if len(value) > 0:
            config[key] = value[0]


def init_dict_if_absent(config: dict, key: str):
    if key not in config or type(config[key]) != dict:
        config[key] = {}
