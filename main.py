import gzip
import os
import platform
import subprocess
import sys
from tkinter import filedialog, messagebox

from docx import Document

from config import *
from songbook2docx.parser import parse_html
from songbook2docx.song import Song
from songbook2docx.utils import style_manager
from songbook2docx.utils import text_width_provider


def exit_with_message(message: str, code: int):
    messagebox.showerror("Błąd przy generowaniu śpiewnika", message)
    sys.exit(code)


# load file
input_file: str
if len(sys.argv) < 2:
    input_file = filedialog.askopenfilename(
        title="Wybierz śpiewnik do konwersji",
        filetypes=[("Śpiewnik Michała Makosia", INPUT_FILE_EXT)],
        initialdir=INITIAL_OPEN_PATH,
        defaultextension=INPUT_FILE_EXT)
    if not input_file:
        exit_with_message("Nie wybrano pliku, z którego ma być wygenerowany śpiewnik", -1)
else:
    input_file = sys.argv[1]

# validate file
file_path = path.dirname(input_file)
file_name, file_extension = path.splitext(path.basename(input_file))
if file_extension != INPUT_FILE_EXT:
    exit_with_message(f"Wybrany plik ma niepoprawny format.\nProgram akceptuje tylko pliki z rozszerzeniem {INPUT_FILE_EXT}", -2)

# exe_path = path.dirname(sys.executable)
if getattr(sys, 'frozen', False):
    exe_path = path.dirname(sys.executable)
else:
    exe_path = path.dirname(__file__)
project_path = path.dirname(__file__)

config = Config(parse_config(path.join(exe_path, "conf.ini")))

with gzip.open(input_file, "r") as file:
    doc: Document = Document(path.join(project_path, TEMPLATE_PATH))

    songs: list[Song] = parse_html(file.read().decode("UTF-8"))

    style_manager.init(doc, config.config)
    text_width_provider.init_fonts(path.join(project_path, "fonts"), config.wanted_font, DEFAULT_FONT, config.font_size)

    content_par = [par for par in doc.paragraphs if "SongbookContent" in par.text][0]

    for song in songs:
        song.apply_flags(config.get_chord_flags())
        song.transpose()
        song.add_paragraphs_to_doc(content_par, config.tab_stops_offset, config.show_author)

    content_el = content_par._element
    content_el.getparent().remove(content_el)
    content_el._p = content_el._element = None

    filename = filedialog.asksaveasfilename(
        title="Zapisz śpiewnik jako",
        filetypes=[("Dokument programu Word", ".docx")],
        initialdir=file_path,
        initialfile=file_name,
        defaultextension=".docx")
    if filename:
        doc.save(filename)

        # open in word or other
        if platform.system() == 'Darwin':  # macOS
            subprocess.call(('open', filename))
        elif platform.system() == 'Windows':  # Windows
            os.startfile(filename)
        else:  # linux variants
            subprocess.call(('xdg-open', filename))
